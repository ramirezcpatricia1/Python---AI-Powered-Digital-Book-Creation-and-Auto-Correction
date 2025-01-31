from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor

doc = Document('Version_con Columna1.docx')
opening_quotes = ['“', '"']
closing_quotes = ['”', '"']

for i, paragraph in enumerate(doc.paragraphs):
     if i > 10:
        #**titulos**
          if paragraph.text.strip().upper().startswith("CAPÍTULO"):
               paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
               for run in paragraph.runs:
                    run.font.size = Pt(16)
                    run.font.bold = True
                    paragraph.paragraph_format.space_before = Pt(24)
                    paragraph.paragraph_format.space_after = Pt(12)
                    run.font.name = 'Montserrat' 
                    
        #**Dialogos**
          elif paragraph.text.strip().startswith("—"):
             if not paragraph.text.strip().startswith("— "):
                  paragraph.text = paragraph.text.replace("—", "— ")
             paragraph_format = paragraph.paragraph_format
             paragraph_format.line_spacing = 1.5
             paragraph_format.first_line_indent = None
             paragraph_format.space_after = Pt(6)
             paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
             for run in paragraph.runs:
                run.font.italic = True
                run.font.name = 'Open Sans Italic'
                run.font.size = Pt(12)

          #**pensamientos**
          elif any(q in paragraph.text for q in opening_quotes):
               full_text = paragraph.text.strip()
               paragraph.clear()

               for opening, closing in zip(opening_quotes, closing_quotes):
                    if opening in full_text and closing in full_text:
                         parts = full_text.split(closing)
                         for part in parts:
                              if opening in part:
                                   subparts = part.split(opening)
                                   # Texto antes de las comillas
                                   if subparts[0]:
                                        run = paragraph.add_run(subparts[0])
                                        run.font.name = 'Open Sans'
                                        run.font.size = Pt(12)
                                   # Texto dentro de las comillas
                                   run = paragraph.add_run(f'{opening}{subparts[1]}{closing}')
                                   run.font.italic = True
                                   run.font.color.rgb = RGBColor(44, 62, 80)
                                   run.font.name = 'Open Sans Italic'
                                   run.font.size = Pt(12)
                                   run.font.bold = True 
                                   # Texto después de las comillas
                              else:
                                   run = paragraph.add_run(part)
                                   run.font.name = 'Open Sans'
                                   run.font.size = Pt(12)
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.line_spacing = 1.5
                    paragraph_format.space_after = Pt(6)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        #**Cuerpo de Texto**
          else:
              paragraph_format = paragraph.paragraph_format
              paragraph_format.line_spacing = 1.5
              paragraph_format.first_line_indent = Cm(0.5)
              paragraph_format.space_after = Pt(6)
              paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

              for run in paragraph.runs:
                   run.font.name = 'Open Sans'
                   run.font.size = Pt(12) 
doc.save('R_Version_con Columna1.docx')